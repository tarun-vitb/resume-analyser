"""
Advanced Resume Text Extraction Module
Handles PDF and DOCX files with intelligent text cleaning and normalization
"""

import re
import logging
from typing import Dict, List, Optional, Tuple
from pathlib import Path
import fitz  # PyMuPDF
from docx import Document
import spacy
from dataclasses import dataclass

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class ExtractedData:
    """Structured data extracted from resume"""
    raw_text: str
    cleaned_text: str
    name: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    skills: List[str] = None
    experience: List[str] = None
    education: List[str] = None
    sections: Dict[str, str] = None

class DocumentProcessor:
    """Advanced document processing with NLP-powered extraction"""
    
    def __init__(self):
        """Initialize with spaCy model for NER"""
        try:
            self.nlp = spacy.load("en_core_web_sm")
        except OSError:
            logger.warning("spaCy model not found. Install with: python -m spacy download en_core_web_sm")
            self.nlp = None
        
        # Common section headers for resume parsing
        self.section_patterns = {
            'experience': r'(?i)(work\s+experience|professional\s+experience|employment|career|work\s+history)',
            'education': r'(?i)(education|academic|qualifications|degrees)',
            'skills': r'(?i)(skills|technical\s+skills|competencies|expertise)',
            'projects': r'(?i)(projects|portfolio|work\s+samples)',
            'certifications': r'(?i)(certifications|certificates|licenses)',
            'summary': r'(?i)(summary|objective|profile|about)'
        }
        
        # Enhanced skill extraction patterns
        self.skill_patterns = [
            # Programming Languages
            r'\b(?:Python|Java|JavaScript|TypeScript|C\+\+|C#|PHP|Ruby|Go|Rust|Swift|Kotlin|Scala|R|MATLAB|Dart)\b',
            
            # Web Technologies
            r'\b(?:React\.?js?|Angular|Vue\.?js?|Node\.?js?|Express|Django|Flask|Spring|Laravel|ASP\.NET|Ruby on Rails|jQuery|Bootstrap|Tailwind CSS|SASS|LESS|GraphQL|REST(?:ful)? API|WebSockets)\b',
            
            # Databases
            r'\b(?:SQL|MySQL|PostgreSQL|MongoDB|Redis|Elasticsearch|Oracle|SQLite|Microsoft SQL Server|DynamoDB|Cassandra|Firebase|Neo4j|Couchbase|HBase|Snowflake|BigQuery|Redshift)\b',
            
            # Cloud & DevOps
            r'\b(?:AWS|Amazon Web Services|Azure|Google Cloud Platform|GCP|Docker|Kubernetes|K8s|Jenkins|Git|GitHub|GitLab|Bitbucket|CI/CD|Terraform|Ansible|Puppet|Chef|Prometheus|Grafana|Splunk|ELK Stack|Kibana|New Relic|Datadog)\b',
            
            # AI/ML/DL
            r'\b(?:Machine Learning|ML|Deep Learning|DL|AI|Artificial Intelligence|NLP|Natural Language Processing|Computer Vision|CV|Data Science|Data Analysis|Predictive Analytics|Neural Networks|CNN|RNN|LSTM|Transformers|BERT|GPT|OpenAI|Hugging Face|LangChain)\b',
            
            # Data Science & Analytics
            r'\b(?:TensorFlow|PyTorch|Scikit-learn|scikit|Pandas|NumPy|Matplotlib|Seaborn|Plotly|D3\.js|Tableau|Power BI|Excel|Google Sheets|Apache Spark|Hadoop|Hive|Pig|Kafka|Airflow|Apache Beam|Databricks)\b',
            
            # Mobile & Desktop
            r'\b(?:React Native|Flutter|Ionic|Xamarin|Android|iOS|Swift|Kotlin|Java|Objective-C|Electron|Qt|WPF|Windows Forms|GTK\+)\b',
            
            # Other Technologies
            r'\b(?:Blockchain|Ethereum|Smart Contracts|Solidity|Web3|dApps|Cybersecurity|Penetration Testing|Ethical Hacking|Network Security|Cryptography|IoT|Internet of Things|AR|VR|Augmented Reality|Virtual Reality|Unity|Unreal Engine|Game Development|Robotics|RPA|UiPath|Blue Prism|Automation Anywhere)\b',
            
            # Soft Skills & Methodologies
            r'\b(?:Agile|Scrum|Kanban|DevOps|DevSecOps|SRE|TDD|Test-Driven Development|BDD|Behavior-Driven Development|CI/CD|Continuous Integration|Continuous Deployment|Microservices|Monolithic|Serverless|REST|SOAP|GraphQL|gRPC|OAuth|JWT|OAuth 2\.0|OIDC|SAML|SSO|MFA|2FA|Zero Trust)\b'
        ]

    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF with improved handling"""
        try:
            doc = fitz.open(file_path)
            text = ""
            
            for page_num in range(doc.page_count):
                page = doc[page_num]
                # Extract text with layout preservation
                page_text = page.get_text("text")
                if page_text.strip():
                    text += page_text + "\n"
                
                # Fallback to OCR-like extraction if text is sparse
                if len(page_text.strip()) < 50:
                    blocks = page.get_text("dict")
                    for block in blocks.get("blocks", []):
                        if "lines" in block:
                            for line in block["lines"]:
                                for span in line["spans"]:
                                    text += span["text"] + " "
                            text += "\n"
            
            doc.close()
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            return ""

    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX files"""
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            return ""

    def clean_and_normalize_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace and special characters
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'[^\w\s@.-]', ' ', text)
        
        # Remove common PDF artifacts
        text = re.sub(r'\b\d{1,2}/\d{1,2}/\d{2,4}\b', ' ', text)  # Dates
        text = re.sub(r'\bPage\s+\d+\b', ' ', text, flags=re.IGNORECASE)
        
        # Normalize line breaks
        text = re.sub(r'\n+', '\n', text)
        
        return text.strip()

    def extract_contact_info(self, text: str) -> Dict[str, Optional[str]]:
        """Extract contact information using regex patterns"""
        contact_info = {
            'email': None,
            'phone': None,
            'name': None
        }
        
        # Email extraction
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        email_match = re.search(email_pattern, text)
        if email_match:
            contact_info['email'] = email_match.group()
        
        # Phone extraction
        phone_pattern = r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
        phone_match = re.search(phone_pattern, text)
        if phone_match:
            contact_info['phone'] = phone_match.group()
        
        # Name extraction (first few lines, excluding common headers)
        lines = text.split('\n')[:5]
        for line in lines:
            line = line.strip()
            if (len(line.split()) >= 2 and 
                not re.search(r'(resume|cv|curriculum|vitae)', line, re.IGNORECASE) and
                not re.search(r'@|phone|email|address', line, re.IGNORECASE)):
                # Simple name validation
                words = line.split()
                if all(word.replace('.', '').isalpha() for word in words[:2]):
                    contact_info['name'] = ' '.join(words[:2])
                    break
        
        return contact_info

    def extract_skills(self, text: str) -> List[str]:
        """Enhanced skill extraction with NLP and pattern matching"""
        if not text:
            return []
            
        found_skills = set()
        
        # 1. First, try to find a dedicated skills section
        skills_section = self.extract_section(text, 'skills')
        text_to_scan = skills_section if skills_section else text
        
        # 2. Extract skills using regex patterns
        for pattern in self.skill_patterns:
            matches = re.finditer(pattern, text_to_scan, re.IGNORECASE)
            for match in matches:
                skill = match.group(0).strip()
                if skill and len(skill) > 2 and len(skill) < 50:  # Reasonable skill length
                    found_skills.add(skill.lower())  # Normalize to lowercase
        
        # 3. Use spaCy for additional skill extraction if available
        if self.nlp:
            doc = self.nlp(text_to_scan)
            for ent in doc.ents:
                if ent.label_ in ['SKILL', 'TECH', 'PRODUCT', 'ORG']:
                    skill = ent.text.strip()
                    if skill and len(skill) > 2 and len(skill) < 50:
                        found_skills.add(skill.lower())
        
        # 4. Clean up and return unique skills (capitalize properly)
        skills_list = []
        for skill in found_skills:
            # Capitalize properly (e.g., "python" -> "Python", "node.js" -> "Node.js")
            if '.' in skill:
                # Handle tech names like "node.js", "asp.net"
                parts = skill.split('.')
                skill = '.'.join(p.capitalize() for p in parts)
            else:
                skill = skill.capitalize()
            skills_list.append(skill)
        
        return sorted(list(set(skills_list)))  # Remove duplicates and sort

    def extract_section(self, text: str, section_name: str) -> Optional[str]:
        """Extract specific section content"""
        if section_name not in self.section_patterns:
            return None
        
        pattern = self.section_patterns[section_name]
        
        # Find section start
        match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
        if not match:
            return None
        
        start_pos = match.end()
        
        # Find next section or end of text
        remaining_text = text[start_pos:]
        
        # Look for next section header
        next_section_pos = len(remaining_text)
        for other_pattern in self.section_patterns.values():
            if other_pattern != pattern:
                next_match = re.search(other_pattern, remaining_text, re.IGNORECASE | re.MULTILINE)
                if next_match and next_match.start() < next_section_pos:
                    next_section_pos = next_match.start()
        
        section_content = remaining_text[:next_section_pos].strip()
        return section_content if section_content else None

    def extract_structured_data(self, file_path: str) -> ExtractedData:
        """Main method to extract structured data from resume"""
        file_path = Path(file_path)
        
        # Determine file type and extract text
        if file_path.suffix.lower() == '.pdf':
            raw_text = self.extract_text_from_pdf(str(file_path))
        elif file_path.suffix.lower() in ['.docx', '.doc']:
            raw_text = self.extract_text_from_docx(str(file_path))
        else:
            raise ValueError(f"Unsupported file format: {file_path.suffix}")
        
        if not raw_text:
            raise ValueError("No text could be extracted from the document")
        
        # Clean and normalize text
        cleaned_text = self.clean_and_normalize_text(raw_text)
        
        # Extract contact information
        contact_info = self.extract_contact_info(cleaned_text)
        
        # Extract skills
        skills = self.extract_skills(cleaned_text)
        
        # Extract sections
        sections = {}
        for section_name in self.section_patterns.keys():
            section_content = self.extract_section(cleaned_text, section_name)
            if section_content:
                sections[section_name] = section_content
        
        # Extract experience and education (simplified)
        experience = []
        education = []
        
        if 'experience' in sections:
            # Split experience by common patterns
            exp_items = re.split(r'\n(?=\w)', sections['experience'])
            experience = [item.strip() for item in exp_items if len(item.strip()) > 20]
        
        if 'education' in sections:
            # Split education by common patterns
            edu_items = re.split(r'\n(?=\w)', sections['education'])
            education = [item.strip() for item in edu_items if len(item.strip()) > 10]
        
        # Ensure skills is always a list, not None
        if skills is None:
            skills = []
        
        return ExtractedData(
            raw_text=raw_text,
            cleaned_text=cleaned_text,
            name=contact_info['name'],
            email=contact_info['email'],
            phone=contact_info['phone'],
            skills=skills or [],  # Ensure it's never None
            experience=experience or [],
            education=education or [],
            sections=sections or {}
        )

    def validate_extraction(self, extracted_data: ExtractedData) -> Dict[str, bool]:
        """Validate the quality of extraction"""
        validation = {
            'has_text': len(extracted_data.cleaned_text) > 100,
            'has_contact': bool(extracted_data.email or extracted_data.phone),
            'has_skills': len(extracted_data.skills) > 0,
            'has_experience': len(extracted_data.experience) > 0,
            'has_education': len(extracted_data.education) > 0
        }
        
        validation['overall_quality'] = sum(validation.values()) >= 3
        
        return validation
